import docx
d = docx.Document('demo.docx')
# print(d.paragraphs)
# print(d.paragraphs[0])
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)
p = d.paragraphs[1]
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)
print(p.runs[0].bold == True)
print(p.runs[1].bold == True)
p.runs[3].underline = True
p.runs[3].text = 'italic and underline'
d.save('demo2.docx')
p.style = 'Title'
d.save('demo3.docx')

d = docx.Document()
d.add_paragraph('Hello this is a paragraph')
d.add_paragraph('This is a another paragraph')
d.save('demo4.docx')
p = d.paragraphs[0]
p.add_run('This is a new run')
p.runs[1].bold = True
d.save('demo5.docx')

def getText(fileName):
    doc = docx.Document(fileName)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('demo.docx'))